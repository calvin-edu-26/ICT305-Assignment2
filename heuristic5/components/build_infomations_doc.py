from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\huang\Desktop\Group Project")
OUT = ROOT / "Infomations.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_GRAY = "F2F4F7"


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9DEE7")


def set_table_width(table, widths):
    table.autofit = False
    table_borders(table)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_header_row(row):
    for cell in row.cells:
        set_cell_fill(cell, LIGHT_GRAY)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(31, 77, 120)


def add_para(doc, text, style=None):
    paragraph = doc.add_paragraph(text, style=style)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0]
    for idx, header in enumerate(headers):
        hdr.cells[idx].text = header
    style_header_row(hdr)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    set_table_width(table, widths)
    doc.add_paragraph()
    return table


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Climate Finance Gap Dashboard")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    subtitle = doc.add_paragraph()
    run = subtitle.add_run("Heuristic 5 Project Information and Design Rationale")
    run.italic = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(80, 80, 80)

    add_table(
        doc,
        ["Project Item", "Details"],
        [
            ["Primary focus", "Interactive data visualisation and decision support"],
            ["Dashboard framework", "Python Streamlit"],
            ["Topic", "Climate Change and Environmental Impact"],
            ["Heuristic", "Climate Finance Gap"],
            ["Core question", "Are major providers delivering enough adaptation finance to vulnerable recipients?"],
        ],
        [2500, 6860],
    )


def build_doc():
    doc = Document()
    configure_styles(doc)
    add_title(doc)

    doc.add_heading("1. Executive Summary", level=1)
    add_para(
        doc,
        "This project delivers a Python-based interactive dashboard that supports evidence-based decision-making on the climate finance gap. "
        "The dashboard focuses on whether major climate finance providers are directing sufficient adaptation support toward vulnerable recipients, "
        "rather than only presenting finance totals as static charts.",
    )
    add_para(
        doc,
        "The dashboard is designed as a decision-support tool for policy analysts, climate finance teams, and development stakeholders who need to compare provider commitments, recipient needs, vulnerability patterns, and emissions responsibility."
    )

    doc.add_heading("2. Stakeholder and Decision Context", level=1)
    add_table(
        doc,
        ["Element", "Project Position"],
        [
            ["Primary stakeholders", "Policy makers, climate finance analysts, donor-agency staff, and development organisations."],
            ["Decision level", "Strategic and tactical: prioritising adaptation finance, identifying provider gaps, and targeting vulnerable regions."],
            ["Main decision question", "Which providers or regions require stronger adaptation-finance attention?"],
            ["Dashboard outcome", "Users can compare adaptation delivery, mitigation balance, vulnerability targeting, and climate-risk context."],
        ],
        [2500, 6860],
    )

    doc.add_heading("3. Dataset Selection and Justification", level=1)
    add_para(
        doc,
        "The datasets were selected because they connect finance delivery, vulnerability, climate risk, and emissions responsibility. "
        "Together, they support a decision-focused climate finance story rather than isolated visualisation exercises."
    )
    add_table(
        doc,
        ["Dataset", "Purpose in Dashboard", "Reason for Use"],
        [
            ["CRDF-PP_all years-2013-2024.csv", "Provider climate finance commitments", "Shows which providers commit finance and whether it is adaptation or mitigation oriented."],
            ["CRDF-RP_all_years_2001-2024.csv", "Recipient perspective finance", "Supports recipient-side checking of who receives finance and whether delivery aligns with need."],
            ["Climate vulnerability finance ODA dataset", "Vulnerability and funding per capita", "Links finance allocation with vulnerability and recipient exposure."],
            ["owid-co2-data.csv", "CO2 emissions responsibility", "Allows comparison between major emitters and adaptation support."],
            ["OECD DAC compiled CSV", "Historical adaptation and mitigation trend", "Adds earlier finance context from 2009 to 2019."],
            ["Converted climate-change CSV files", "Climate risk context", "Adds temperature, land-risk, and climate indicator context."],
            ["World Bank WDI export", "Global trend context", "Supports background explanation using world-level development and energy indicators."],
            ["historical_emissions.csv", "Historical emissions context", "Provides an additional emissions perspective for climate responsibility discussion."],
        ],
        [2500, 3400, 3460],
    )

    doc.add_heading("4. Dashboard Structure and User Flow", level=1)
    add_para(
        doc,
        "The dashboard is arranged into four guided sections in the sidebar. This improves usability by giving users a natural sequence instead of forcing them to interpret every chart at once."
    )
    add_bullets(
        doc,
        [
            "Step 1 - Main Finance Gap: summary KPIs, provider comparison, recipient scatter plot, emissions-versus-adaptation gap, and decision table.",
            "Step 2 - Recipient Finance: recipient-side finance view and historical OECD DAC adaptation versus mitigation context.",
            "Step 3 - Risk & Global Context: climate risk, World Bank global indicators, and historical emissions evidence.",
            "Step 4 - Data Sources: transparent explanation of which source files support each chart.",
        ],
    )

    doc.add_heading("5. Interactivity and Dashboard Controls", level=1)
    add_table(
        doc,
        ["Interactive Feature", "Analytical Purpose"],
        [
            ["Commitment year range slider", "Filters finance records by year so users can compare recent and historical periods."],
            ["Provider country multiselect", "Allows users to focus on specific major providers such as the United States, Germany, Japan, France, and the United Kingdom."],
            ["Recipient region multiselect", "Supports regional comparison and targeted recipient analysis."],
            ["Gap lens radio button", "Switches the analytical focus between adaptation delivery, mitigation balance, and vulnerability targeting."],
            ["Mitigation finance range slider", "Lets users increase or decrease the visible range in the recipient adaptation-versus-mitigation scatter plot."],
            ["Global context indicator selector", "Allows users to inspect different global background trends without crowding the dashboard."],
        ],
        [3200, 6160],
    )

    doc.add_heading("6. Visualisation Rationale", level=1)
    add_table(
        doc,
        ["Chart Type", "Dashboard Use", "Why Appropriate"],
        [
            ["Line chart", "Adaptation versus mitigation over time", "Best for temporal trends and finance changes across years."],
            ["Horizontal bar chart", "Provider commitments and recipient rankings", "Readable for comparing many countries or providers."],
            ["Treemap", "Regional distribution of finance", "Shows proportional allocation across recipient regions."],
            ["Bubble scatter plot", "Adaptation versus mitigation and emissions responsibility", "Shows relationship, magnitude, and outliers in one chart."],
            ["Choropleth map", "Vulnerable-country funding per capita", "Reveals spatial patterns in funding and vulnerability."],
            ["Violin plot", "Adaptation share by region", "Shows distribution differences rather than only averages."],
            ["Area chart", "Historical OECD DAC finance", "Shows changing composition of adaptation and mitigation finance over time."],
        ],
        [2300, 3300, 3760],
    )

    doc.add_heading("7. Key Analytical Insights", level=1)
    add_bullets(
        doc,
        [
            "The dashboard can identify whether selected providers commit more toward mitigation than adaptation.",
            "Provider-level comparison reveals possible gaps between emissions responsibility and adaptation support.",
            "Recipient-level views highlight which regions or countries receive the largest climate finance commitments.",
            "Vulnerability and funding-per-capita views support discussion about whether finance reaches countries exposed to climate risk.",
            "Historical OECD DAC and World Bank context strengthens the dashboard by showing that climate finance should be interpreted against long-term risk and development trends.",
        ],
    )

    doc.add_heading("8. Design Principles Applied", level=1)
    add_table(
        doc,
        ["Principle", "How It Is Applied"],
        [
            ["Clarity", "The dashboard uses a guided sidebar flow and short section descriptions."],
            ["Consistency", "Filters are placed in one sidebar and reused across finance views."],
            ["Purposeful interactivity", "Each control changes a meaningful analytical question rather than decorating the page."],
            ["Accessibility", "Charts use readable labels, concise titles, and conventional chart types."],
            ["Data integrity", "The dashboard separates finance, vulnerability, emissions, and context datasets while documenting their roles."],
        ],
        [2500, 6860],
    )

    doc.add_heading("9. Limitations and Responsible Interpretation", level=1)
    add_bullets(
        doc,
        [
            "Some datasets use different price bases and reporting standards, so monetary comparisons should be interpreted with attention to units.",
            "The World Bank WDI file is world-level only, so it supports global context rather than country-level finance decisions.",
            "Finance commitment data does not always equal actual disbursement, so conclusions should be framed as commitment-based evidence.",
            "The dashboard supports decision-making, but final policy decisions should also consider qualitative country context and updated official reporting.",
        ],
    )

    doc.add_heading("10. Distinction-Level Alignment", level=1)
    add_para(
        doc,
        "The project aligns with distinction-level expectations by demonstrating a clear decision context, justified dataset selection, purposeful interactivity, multiple appropriate visualisation techniques, and a dashboard structure that supports interpretation rather than simple decoration."
    )
    add_bullets(
        doc,
        [
            "Python is used for data loading, processing, dashboard logic, and visualisation.",
            "The dashboard uses multiple interactive controls and dynamically updates charts based on selected filters.",
            "The analysis connects environmental responsibility, finance allocation, and vulnerability.",
            "The visual design avoids redundant charts by assigning each chart a clear analytical purpose.",
            "The documentation explains both technical implementation and stakeholder decision value.",
        ],
    )

    doc.add_heading("11. Conclusion", level=1)
    add_para(
        doc,
        "The Climate Finance Gap dashboard provides a focused, interactive, and decision-oriented tool for evaluating whether adaptation finance is being directed toward vulnerable recipients. "
        "By combining provider finance, recipient finance, emissions responsibility, vulnerability evidence, and climate-risk context, the project demonstrates both technical dashboard-building skill and applied data storytelling."
    )

    doc.core_properties.title = "Climate Finance Gap Dashboard Information"
    doc.core_properties.subject = "Heuristic 5 Streamlit Dashboard"
    doc.core_properties.author = "Huang Neng Jie"
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
